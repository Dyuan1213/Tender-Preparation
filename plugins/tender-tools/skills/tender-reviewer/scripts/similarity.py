#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
标书查重脚本（tender-reviewer 配套）

对两份及以上 .docx 标书做两两比对，输出：
- 文本相似度：Jaccard（5-字 shingles，适合中文）+ SequenceMatcher.ratio()
- 段落雷同：完全相同段落占比、相同段落数
- docx 元数据：author、lastModifiedBy、created、modified、revision
- 段落样式集合及其重合度
- 风险等级：high / medium / low（按文本相似度、段落雷同、元数据综合判定）

用法：
    python similarity.py --files A.docx B.docx [C.docx ...] [--output report.json] [--text-cap 100000]

依赖：python-docx + 标准库；无其他第三方依赖。
"""

from __future__ import annotations

import argparse
import json
import re
import sys
from datetime import datetime, timedelta
from difflib import SequenceMatcher
from itertools import combinations
from pathlib import Path
from typing import Optional


def _docx() -> "module":
    try:
        from docx import Document  # noqa: F401
        import docx
        return docx
    except ImportError:
        print("python-docx 未安装，请运行: pip install python-docx", file=sys.stderr)
        sys.exit(2)


# ---------- 文档读取 ----------

def extract_doc(path: Path) -> dict:
    """读取一份 docx，返回 {paragraphs, tables_text, all_text, styles, props}"""
    docx = _docx()
    d = docx.Document(str(path))

    paragraphs: list[str] = []
    styles: list[str] = []
    for p in d.paragraphs:
        t = p.text.strip()
        if t:
            paragraphs.append(t)
            styles.append((p.style.name if p.style else "") or "")

    tables_text: list[str] = []
    for tb in d.tables:
        for row in tb.rows:
            for c in row.cells:
                t = c.text.strip()
                if t:
                    tables_text.append(t)

    all_text = "\n".join(paragraphs + tables_text)

    cp = d.core_properties
    props = {
        "author": getattr(cp, "author", None) or "",
        "last_modified_by": getattr(cp, "last_modified_by", None) or "",
        "created": _iso(getattr(cp, "created", None)),
        "modified": _iso(getattr(cp, "modified", None)),
        "revision": getattr(cp, "revision", None),
        "title": getattr(cp, "title", None) or "",
    }

    return {
        "path": str(path),
        "name": path.name,
        "paragraphs": paragraphs,
        "tables_text": tables_text,
        "all_text": all_text,
        "styles": styles,
        "n_paragraphs": len(paragraphs),
        "n_tables": len(d.tables),
        "n_sections": len(d.sections),
        "props": props,
    }


def _iso(dt) -> Optional[str]:
    if dt is None:
        return None
    try:
        return dt.isoformat()
    except Exception:
        return str(dt)


# ---------- 文本规整与相似度 ----------

_WS_RE = re.compile(r"\s+")
_PUNCT_RE = re.compile(r"[\s，。；：、！？“”‘’（）【】《》—…·\-\.,;:!?\"'()\[\]<>]")

def normalize_for_para(text: str) -> str:
    """段落级规整：去多余空白、保留标点，用于"完全相同段落"判定。"""
    return _WS_RE.sub(" ", text).strip()


def normalize_for_shingle(text: str) -> str:
    """字符 shingle 用：去空白与常见标点，避免标点差异影响相似度。"""
    return _PUNCT_RE.sub("", text)


def char_shingles(text: str, n: int = 5) -> set[str]:
    """5-字符 shingles（适合中文与无空格分词的中英混排）"""
    t = normalize_for_shingle(text)
    if len(t) < n:
        return set([t]) if t else set()
    return {t[i:i + n] for i in range(len(t) - n + 1)}


def jaccard(a: set, b: set) -> float:
    if not a and not b:
        return 0.0
    inter = len(a & b)
    union = len(a | b)
    return inter / union if union else 0.0


def seq_ratio(a: str, b: str, cap: int = 100_000) -> float:
    """SequenceMatcher.ratio，超长截断以保证时延可控。"""
    a2 = a[:cap]
    b2 = b[:cap]
    sm = SequenceMatcher(a=a2, b=b2, autojunk=False)
    return sm.ratio()


def identical_para_stats(p1: list[str], p2: list[str]) -> dict:
    s1 = {normalize_for_para(x) for x in p1 if normalize_for_para(x)}
    s2 = {normalize_for_para(x) for x in p2 if normalize_for_para(x)}
    if not s1 or not s2:
        return {"identical_count": 0, "ratio_min": 0.0, "ratio_avg": 0.0}
    common = s1 & s2
    ratio_min = len(common) / min(len(s1), len(s2))
    ratio_avg = len(common) / ((len(s1) + len(s2)) / 2)
    return {
        "identical_count": len(common),
        "ratio_min": round(ratio_min, 4),
        "ratio_avg": round(ratio_avg, 4),
        "examples": sorted(common, key=len, reverse=True)[:5],
    }


def styles_overlap(s1: list[str], s2: list[str]) -> dict:
    set1, set2 = set(s1), set(s2)
    if not set1 and not set2:
        return {"jaccard": 0.0, "common": []}
    return {
        "jaccard": round(jaccard(set1, set2), 4),
        "common": sorted(set1 & set2),
    }


# ---------- 元数据匹配 ----------

def meta_match(p1: dict, p2: dict) -> dict:
    matches: dict = {}
    same_author = bool(p1.get("author")) and p1.get("author") == p2.get("author")
    same_lmb = bool(p1.get("last_modified_by")) and p1.get("last_modified_by") == p2.get("last_modified_by")
    matches["same_author"] = same_author
    matches["same_last_modified_by"] = same_lmb
    matches["author_a"] = p1.get("author", "")
    matches["author_b"] = p2.get("author", "")
    matches["last_modified_by_a"] = p1.get("last_modified_by", "")
    matches["last_modified_by_b"] = p2.get("last_modified_by", "")

    matches["created_close"] = False
    matches["modified_close"] = False
    matches["created_delta_seconds"] = None
    matches["modified_delta_seconds"] = None
    for key, flag in (("created", "created_close"), ("modified", "modified_close")):
        d1 = _parse_iso(p1.get(key))
        d2 = _parse_iso(p2.get(key))
        if d1 and d2:
            delta = abs((d1 - d2).total_seconds())
            matches[f"{key}_delta_seconds"] = int(delta)
            matches[flag] = delta < 300  # 5 分钟内视为"几乎一致"
    return matches


def _parse_iso(s: Optional[str]) -> Optional[datetime]:
    if not s:
        return None
    try:
        return datetime.fromisoformat(s.replace("Z", "+00:00"))
    except Exception:
        return None


# ---------- 风险等级 ----------

def risk_level(pair: dict) -> dict:
    reasons: list[str] = []
    level = "low"

    jc = pair["text"]["jaccard_5gram"]
    sr = pair["text"]["seq_ratio"]
    ip = pair["paragraphs"]["ratio_min"]
    mm = pair["metadata"]

    # 高风险触发
    if jc >= 0.6:
        level = "high"; reasons.append(f"文本 Jaccard={jc:.2f} ≥ 0.6")
    if ip >= 0.5:
        level = "high"; reasons.append(f"相同段落占比={ip:.0%} ≥ 50%")
    if mm.get("same_author"):
        level = "high"; reasons.append(f"作者相同：{mm.get('author_a')!r}")
    if mm.get("same_last_modified_by"):
        level = "high"; reasons.append(f"最后修改人相同：{mm.get('last_modified_by_a')!r}")
    if mm.get("created_close"):
        level = "high"; reasons.append(f"创建时间几乎一致（差 {mm['created_delta_seconds']}s）")

    # 中风险（如未已被判高）
    if level != "high":
        if 0.3 <= jc < 0.6:
            level = "medium"; reasons.append(f"文本 Jaccard={jc:.2f} ∈ [0.3, 0.6)")
        if 0.2 <= ip < 0.5:
            level = "medium"; reasons.append(f"相同段落占比={ip:.0%} ∈ [20%, 50%)")
        if sr >= 0.5:
            level = "medium"; reasons.append(f"SequenceMatcher 比={sr:.2f} ≥ 0.5")

    if not reasons:
        reasons.append("文本/段落/元数据相似度均偏低，未触发风险阈值")

    return {"level": level, "reasons": reasons}


# ---------- 主流程 ----------

def compare_pair(a: dict, b: dict, text_cap: int = 100_000) -> dict:
    sh_a = char_shingles(a["all_text"], n=5)
    sh_b = char_shingles(b["all_text"], n=5)
    jc = round(jaccard(sh_a, sh_b), 4)
    sr = round(seq_ratio(a["all_text"], b["all_text"], cap=text_cap), 4)

    para = identical_para_stats(a["paragraphs"], b["paragraphs"])
    sty = styles_overlap(a["styles"], b["styles"])
    mm = meta_match(a["props"], b["props"])

    pair = {
        "a": a["name"],
        "b": b["name"],
        "structure": {
            "paragraphs_a": a["n_paragraphs"],
            "paragraphs_b": b["n_paragraphs"],
            "tables_a": a["n_tables"],
            "tables_b": b["n_tables"],
            "sections_a": a["n_sections"],
            "sections_b": b["n_sections"],
        },
        "text": {"jaccard_5gram": jc, "seq_ratio": sr},
        "paragraphs": para,
        "styles": sty,
        "metadata": mm,
    }
    pair["risk"] = risk_level(pair)
    return pair


def main():
    ap = argparse.ArgumentParser(description="标书查重（tender-reviewer 配套）")
    ap.add_argument("--files", nargs="+", required=True, help="2 个或多个 .docx 路径")
    ap.add_argument("--output", default=None, help="输出 JSON 路径（缺省则打印到 stdout）")
    ap.add_argument("--text-cap", type=int, default=100_000,
                    help="SequenceMatcher 全文截断长度（字符），默认 100000")
    args = ap.parse_args()

    paths = [Path(p) for p in args.files]
    for p in paths:
        if not p.exists():
            print(f"文件不存在：{p}", file=sys.stderr); sys.exit(2)
    if len(paths) < 2:
        print("至少需要 2 个文件", file=sys.stderr); sys.exit(2)

    docs = [extract_doc(p) for p in paths]

    pairs = [compare_pair(a, b, text_cap=args.text_cap) for a, b in combinations(docs, 2)]

    report = {
        "files": [{"name": d["name"], "path": d["path"], "props": d["props"],
                   "n_paragraphs": d["n_paragraphs"], "n_tables": d["n_tables"]} for d in docs],
        "pairs": pairs,
        "summary": {
            "n_files": len(docs),
            "n_pairs": len(pairs),
            "high_risk_pairs": sum(1 for p in pairs if p["risk"]["level"] == "high"),
            "medium_risk_pairs": sum(1 for p in pairs if p["risk"]["level"] == "medium"),
            "low_risk_pairs": sum(1 for p in pairs if p["risk"]["level"] == "low"),
        },
    }

    out = json.dumps(report, ensure_ascii=False, indent=2)
    if args.output:
        Path(args.output).parent.mkdir(parents=True, exist_ok=True)
        Path(args.output).write_text(out, encoding="utf-8")
        print(f"查重报告已生成：{args.output}")
        # 同时打印精简摘要到 stdout 便于 AI 读取
        summary = {
            "summary": report["summary"],
            "pairs": [
                {"a": p["a"], "b": p["b"],
                 "jaccard": p["text"]["jaccard_5gram"],
                 "seq_ratio": p["text"]["seq_ratio"],
                 "identical_para_ratio": p["paragraphs"]["ratio_min"],
                 "risk": p["risk"]["level"],
                 "reasons": p["risk"]["reasons"]}
                for p in pairs
            ],
        }
        print(json.dumps(summary, ensure_ascii=False, indent=2))
    else:
        print(out)


if __name__ == "__main__":
    main()
