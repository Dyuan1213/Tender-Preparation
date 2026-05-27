#!/usr/bin/env python3
"""
标书 Word 文档生成脚本
生成格式：宋体、有页码、无页眉、无页脚装饰内容

用法：
  python create_tender_docx.py --content-file content.json --output output.docx [--title "投标文件"]

content.json 格式：
{
  "title": "投标文件",
  "project_name": "xxx项目",
  "sections": [
    {
      "level": 1,
      "heading": "第一章 投标函",
      "content": "正文内容...",
      "type": "fixed_format"   // fixed_format | substantive | free_writing
    },
    {
      "level": 2,
      "heading": "1.1 投标函正文",
      "content": "..."
    }
  ]
}
"""

import argparse
import json
import os
import re
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.section import WD_SECTION
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import docx.opc.constants
except ImportError:
    print("ERROR: python-docx 未安装，请运行: pip install python-docx", file=sys.stderr)
    sys.exit(1)


SONG_TI = "宋体"

# 正文内有序标题的层级识别（相对于所在章节再向下逐级递进）
# depth 0 最高，依次递增；最终大纲级别 = 章节 level + depth
_CN_NUM = "一二三四五六七八九十百零两"
_HEADING_PATTERNS = [
    (re.compile(r"^[" + _CN_NUM + r"]+、"), 0),        # 一、二、三、
    (re.compile(r"^（[" + _CN_NUM + r"]+）"), 1),       # （一）（二）
    (re.compile(r"^[0-9]{1,2}[\.．](?![0-9])"), 2),     # 1. 2.（排除 1.5 这类小数）
    (re.compile(r"^（[0-9]{1,2}）"), 3),                # （1）（2）
]


def detect_heading_depth(line):
    """识别正文内的有序标题，返回相对深度（0 最高），非标题返回 None"""
    s = line.strip()
    if not s or len(s) > 50:
        return None
    # 句末带句号/分号/逗号的视为正文句子，不当作标题
    if s.endswith(("。", "；", "，")):
        return None
    for pat, depth in _HEADING_PATTERNS:
        if pat.match(s):
            return depth
    return None


def set_font(run, size_pt, bold=False):
    """设置字体为宋体"""
    run.font.name = SONG_TI
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    # 中文字体需要同时设置 rPr 中的 rFonts
    r = run._r
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:eastAsia"), SONG_TI)
    rFonts.set(qn("w:ascii"), SONG_TI)
    rFonts.set(qn("w:hAnsi"), SONG_TI)
    rPr.insert(0, rFonts)


def set_paragraph_font(paragraph, size_pt, bold=False):
    """设置段落中所有 run 的字体"""
    for run in paragraph.runs:
        set_font(run, size_pt, bold)
    # 同时设置段落默认字体
    pPr = paragraph._p.get_or_add_pPr()
    rPr = OxmlElement("w:rPr")
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:eastAsia"), SONG_TI)
    rFonts.set(qn("w:ascii"), SONG_TI)
    rFonts.set(qn("w:hAnsi"), SONG_TI)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(int(size_pt * 2)))
    szCs = OxmlElement("w:szCs")
    szCs.set(qn("w:val"), str(int(size_pt * 2)))
    rPr.append(rFonts)
    rPr.append(sz)
    rPr.append(szCs)
    if bold:
        b = OxmlElement("w:b")
        rPr.append(b)
    pPr.append(rPr)


def add_page_number(doc):
    """在页脚添加居中页码：第 X 页 / 共 Y 页"""
    section = doc.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False

    # 清空现有页脚内容
    for para in footer.paragraphs:
        for run in para.runs:
            run.text = ""

    # 使用第一个段落（如果没有则添加）
    if footer.paragraphs:
        p = footer.paragraphs[0]
    else:
        p = footer.add_paragraph()

    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 清空段落内容
    p.clear()

    # 构建 "第 {页码} 页" 字段
    run = p.add_run("第 ")
    set_font(run, 10)

    # 插入页码域
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = " PAGE "
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")

    run2 = p.add_run()
    set_font(run2, 10)
    run2._r.append(fldChar1)
    run2._r.append(instrText)
    run2._r.append(fldChar2)

    run3 = p.add_run(" 页")
    set_font(run3, 10)


def add_cover_page(doc, title, project_name=""):
    """添加封面页"""
    # 封面大标题
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    cover_title = doc.add_paragraph()
    cover_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cover_title.add_run(title)
    set_font(run, 22, bold=True)

    doc.add_paragraph()

    if project_name:
        proj_para = doc.add_paragraph()
        proj_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run2 = proj_para.add_run(f"项目名称：{project_name}")
        set_font(run2, 14)

    doc.add_paragraph()
    doc.add_paragraph()

    # 提示填写区域
    info_items = [
        "投标人（盖章）：【请填写：投标人全称】",
        "法定代表人（签字）：【请填写】",
        "日期：          年     月     日",
    ]
    for item in info_items:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(item)
        set_font(run, 12)

    # 封面后分页
    doc.add_page_break()


def set_document_default_font(doc):
    """设置文档默认字体为宋体"""
    styles = doc.styles
    # 设置默认段落字体
    style = styles["Normal"]
    style.font.name = SONG_TI
    style.font.size = Pt(12)
    # 设置中文字体
    element = style.element
    rPr = element.find(qn("w:rPr"))
    if rPr is None:
        rPr = OxmlElement("w:rPr")
        element.append(rPr)
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), SONG_TI)
    rFonts.set(qn("w:ascii"), SONG_TI)
    rFonts.set(qn("w:hAnsi"), SONG_TI)


def process_content_text(doc, text, base_font_size=12, base_level=1):
    """
    处理正文文本，支持简单的 Markdown 类格式：
    - 空行分段
    - | 开头的行作为表格行（用 | 分隔列）
    - 【...】标注为高亮提示
    - 一、/（一）/1./（1）等有序标题赋予对应大纲级别

    base_level：所在章节的标题层级，正文有序标题在此基础上逐级递进。
    """
    if not text:
        return

    lines = text.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()

        # 检测图片块：[[IMG:绝对路径|图注]]，居中插入，下方居中灰色图注
        s_img = line.strip()
        if s_img.startswith("[[IMG:") and s_img.endswith("]]"):
            inner = s_img[6:-2]
            img_path, caption = (inner.split("|", 1) + [""])[:2]
            img_path = img_path.strip()
            if os.path.exists(img_path):
                pic_p = doc.add_paragraph()
                pic_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                pic_p.paragraph_format.first_line_indent = Pt(0)
                run = pic_p.add_run()
                try:
                    run.add_picture(img_path, width=Cm(13))
                except Exception:
                    run.add_picture(img_path)
                if caption.strip():
                    cap_p = doc.add_paragraph()
                    cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    cap_p.paragraph_format.first_line_indent = Pt(0)
                    cr = cap_p.add_run(caption.strip())
                    set_font(cr, 10.5)
                    cr.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
            i += 1
            continue

        # 检测表格块（连续的 | 开头行）
        if line.startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i].strip())
                i += 1

            # 解析并添加表格
            add_simple_table(doc, table_lines, base_font_size)
            continue

        # 检测正文内有序标题（一、/（一）/1./（1）），赋予大纲级别
        depth = detect_heading_depth(line) if line else None
        if depth is not None:
            hp = doc.add_paragraph()
            run = hp.add_run(line.strip())
            set_font(run, base_font_size, bold=True)

            hpPr = hp._p.get_or_add_pPr()
            outline = OxmlElement("w:outlineLvl")
            outline.set(qn("w:val"), str(min(base_level + depth, 8)))
            hpPr.append(outline)

            # 标题不首行缩进、段前段后不加空行
            hp.paragraph_format.first_line_indent = Pt(0)
            hp.paragraph_format.space_before = Pt(0)
            hp.paragraph_format.space_after = Pt(0)
            i += 1
            continue

        # 普通段落
        if line:
            p = doc.add_paragraph()
            # 首行缩进2字符：用 firstLineChars（单位为1/100字符），随字号自适应
            pPr = p._p.get_or_add_pPr()
            ind = OxmlElement("w:ind")
            ind.set(qn("w:firstLineChars"), "200")  # 2 个字符
            ind.set(qn("w:firstLine"), str(int(base_font_size * 2 * 20)))  # 兼容回退值
            pPr.append(ind)

            # 处理【】标注
            parts = line.split("【")
            first_part = parts[0]
            if first_part:
                run = p.add_run(first_part)
                set_font(run, base_font_size)

            for part in parts[1:]:
                if "】" in part:
                    bracket_content, rest = part.split("】", 1)
                    # 高亮显示【...】内容
                    bracket_run = p.add_run("【" + bracket_content + "】")
                    set_font(bracket_run, base_font_size)
                    bracket_run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)  # 红色
                    bracket_run.font.bold = True
                    if rest:
                        rest_run = p.add_run(rest)
                        set_font(rest_run, base_font_size)
                else:
                    run = p.add_run("【" + part)
                    set_font(run, base_font_size)
        # 空行仅作段落分隔，不插入空段落（避免正文段间出现空行）

        i += 1


def add_simple_table(doc, table_lines, font_size=11):
    """从 Markdown 表格行创建 Word 表格"""
    # 过滤掉分隔行（---|---）
    data_lines = [
        line for line in table_lines
        if not all(c in "|-: " for c in line)
    ]
    if not data_lines:
        return

    rows_data = []
    for line in data_lines:
        cells = [c.strip() for c in line.strip("|").split("|")]
        rows_data.append(cells)

    if not rows_data:
        return

    num_cols = max(len(row) for row in rows_data)
    table = doc.add_table(rows=len(rows_data), cols=num_cols)
    table.style = "Table Grid"

    for r_idx, row_data in enumerate(rows_data):
        row = table.rows[r_idx]
        for c_idx, cell_text in enumerate(row_data):
            if c_idx < num_cols:
                cell = row.cells[c_idx]
                cell.text = ""
                p = cell.paragraphs[0]
                run = p.add_run(cell_text)
                set_font(run, font_size, bold=(r_idx == 0))

    doc.add_paragraph()  # 表格后空行


def add_section(doc, section_data, section_number=None):
    """添加一个章节"""
    level = section_data.get("level", 1)
    heading_text = section_data.get("heading", "")
    content = section_data.get("content", "")
    section_type = section_data.get("type", "free_writing")

    # 标题统一四号（14pt）加粗，等级通过编号体系区分
    if level == 1:
        font_size = 14
        heading_style = "Heading 1"
    elif level == 2:
        font_size = 14
        heading_style = "Heading 2"
    else:
        font_size = 14
        heading_style = "Heading 3"

    # 一级章节前分页
    if level == 1 and section_number and section_number > 1:
        doc.add_page_break()

    # 添加标题：套用内置标题样式，使 Word 能识别大纲层级（导航窗格/目录）
    p = doc.add_paragraph(style=heading_style)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(heading_text)
    set_font(run, font_size, bold=True)
    run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)  # 覆盖标题样式默认颜色为黑色

    # 显式设置大纲级别，确保层级被正确识别（level1→0, level2→1, level3→2）
    pPr = p._p.get_or_add_pPr()
    outline = OxmlElement("w:outlineLvl")
    outline.set(qn("w:val"), str(level - 1))
    pPr.append(outline)

    # 标题无首行缩进、段前段后不加空行
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)

    # 类型标注（仅在内容中不显示，用于内部控制）
    # 固定格式章节加说明
    if section_type == "fixed_format" and level == 1:
        note_p = doc.add_paragraph()
        note_run = note_p.add_run("【说明：本章为固定格式，请按招标文件原格式填写标注处内容】")
        set_font(note_run, 10)
        note_run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
        note_run.font.italic = True

    # 添加正文内容
    if content:
        process_content_text(doc, content, base_font_size=12, base_level=level)

    # 设置行距（全局在段落级别控制）
    # 正文行距1.5倍已在段落格式中设置


def create_document(content_data, output_path):
    """创建完整的 Word 文档"""
    doc = Document()

    # 设置页面边距
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

    # 设置默认字体
    set_document_default_font(doc)

    # 设置全局行距：1.5 倍，段前段后不加空行
    from docx.enum.text import WD_LINE_SPACING
    style = doc.styles["Normal"]
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)

    # 清除页眉（确保无内容）
    for sec in doc.sections:
        sec.header.is_linked_to_previous = False
        header = sec.header
        for para in header.paragraphs:
            for run in para.runs:
                run.text = ""

    # 添加封面
    title = content_data.get("title", "投标文件")
    project_name = content_data.get("project_name", "")
    add_cover_page(doc, title, project_name)

    # 添加页码到页脚
    add_page_number(doc)

    # 逐章节添加内容
    sections = content_data.get("sections", [])
    chapter_count = 0
    for sec in sections:
        if sec.get("level", 1) == 1:
            chapter_count += 1
        add_section(doc, sec, chapter_count if sec.get("level", 1) == 1 else None)

    # 保存文档
    doc.save(output_path)
    print(f"文档已生成：{output_path}")


def main():
    parser = argparse.ArgumentParser(description="生成标书 Word 文档")
    parser.add_argument("--content-file", required=True, help="内容 JSON 文件路径")
    parser.add_argument("--output", required=True, help="输出 .docx 文件路径")
    parser.add_argument("--title", default="投标文件", help="文档标题（可被 JSON 中的 title 覆盖）")
    args = parser.parse_args()

    content_file = Path(args.content_file)
    if not content_file.exists():
        print(f"ERROR: 内容文件不存在：{content_file}", file=sys.stderr)
        sys.exit(1)

    with open(content_file, "r", encoding="utf-8") as f:
        content_data = json.load(f)

    # 命令行 title 作为默认值，JSON 中的 title 优先
    if "title" not in content_data:
        content_data["title"] = args.title

    output_path = Path(args.output)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    create_document(content_data, str(output_path))


if __name__ == "__main__":
    main()
