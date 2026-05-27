#!/usr/bin/env python3
"""
Upload a Markdown tender analysis report to Feishu as a document.

Flow:
  1. Convert Markdown report to DOCX
  2. Upload DOCX to Feishu Drive
  3. Import DOCX as a Feishu Doc
  4. Return the document URL

Usage:
    python scripts/feishu_upload.py <report.md> [--folder-token FOLDER_TOKEN]

Credentials (set as environment variables or in .env):
    FEISHU_APP_ID       Your Feishu app's App ID
    FEISHU_APP_SECRET   Your Feishu app's App Secret
    FEISHU_FOLDER_TOKEN (Optional) Target folder token; defaults to root My Drive
"""

import argparse
import json
import os
import sys
import time
from pathlib import Path

BASE_URL = "https://open.feishu.cn/open-apis"


# ── Helpers ──────────────────────────────────────────────────────────────────

def load_credentials():
    """Load app_id and app_secret from env or .env file in the skill root."""
    app_id = os.environ.get("FEISHU_APP_ID")
    app_secret = os.environ.get("FEISHU_APP_SECRET")

    if not app_id or not app_secret:
        # Try .env file in skill directory
        env_path = Path(__file__).parent.parent / ".env"
        if env_path.exists():
            for line in env_path.read_text(encoding="utf-8").splitlines():
                line = line.strip()
                if line.startswith("FEISHU_APP_ID="):
                    app_id = line.split("=", 1)[1].strip().strip('"')
                elif line.startswith("FEISHU_APP_SECRET="):
                    app_secret = line.split("=", 1)[1].strip().strip('"')

    if not app_id or not app_secret:
        print("❌ 未找到飞书凭证。请设置环境变量 FEISHU_APP_ID 和 FEISHU_APP_SECRET，")
        print("   或在 skill 根目录创建 .env 文件（格式见 references/feishu-setup.md）")
        sys.exit(1)

    return app_id, app_secret


def get_tenant_token(app_id: str, app_secret: str) -> str:
    """Get tenant_access_token using app credentials."""
    import urllib.request

    payload = json.dumps({"app_id": app_id, "app_secret": app_secret}).encode()
    req = urllib.request.Request(
        f"{BASE_URL}/auth/v3/tenant_access_token/internal",
        data=payload,
        headers={"Content-Type": "application/json"},
    )
    with urllib.request.urlopen(req) as resp:
        data = json.loads(resp.read())

    if data.get("code") != 0:
        print(f"❌ 获取 token 失败：{data}")
        sys.exit(1)

    return data["tenant_access_token"]


# ── DOCX conversion ───────────────────────────────────────────────────────────

def markdown_to_docx(md_path: Path, docx_path: Path):
    """Convert a Markdown file to DOCX using python-docx with basic formatting."""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        print("⚠️  python-docx 未安装，尝试 pip install python-docx ...")
        os.system("pip install python-docx -q")
        from docx import Document
        from docx.shared import Pt, RGBColor

    doc = Document()

    # Minimal style tweaks
    style = doc.styles["Normal"]
    style.font.name = "微软雅黑"
    style.font.size = Pt(10.5)

    lines = md_path.read_text(encoding="utf-8").splitlines()
    in_table = False
    table_rows = []

    def flush_table():
        nonlocal in_table, table_rows
        if not table_rows:
            return
        # Remove separator row (---|---|---)
        data_rows = [r for r in table_rows if not all(c.strip().replace("-", "").replace("|", "") == "" for c in r)]
        if not data_rows:
            table_rows = []
            in_table = False
            return
        col_count = max(len(r) for r in data_rows)
        t = doc.add_table(rows=len(data_rows), cols=col_count)
        t.style = "Table Grid"
        for ri, row in enumerate(data_rows):
            for ci, cell_text in enumerate(row[:col_count]):
                cell = t.cell(ri, ci)
                cell.text = cell_text.strip()
                if ri == 0:
                    for run in cell.paragraphs[0].runs:
                        run.bold = True
        doc.add_paragraph()
        table_rows = []
        in_table = False

    def parse_table_row(line: str):
        cells = [c for c in line.strip().strip("|").split("|")]
        return cells

    for line in lines:
        # Table detection
        if "|" in line and line.strip().startswith("|"):
            if not in_table:
                in_table = True
            table_rows.append(parse_table_row(line))
            continue
        elif in_table:
            flush_table()

        stripped = line.rstrip()

        # Headings
        if stripped.startswith("#### "):
            doc.add_heading(stripped[5:], level=4)
        elif stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        # Horizontal rule
        elif stripped.startswith("---"):
            doc.add_paragraph("─" * 40)
        # Checkboxes / bullets
        elif stripped.startswith("- [ ] ") or stripped.startswith("- [x] "):
            checked = stripped.startswith("- [x] ")
            text = ("☑ " if checked else "☐ ") + stripped[6:]
            doc.add_paragraph(text, style="List Bullet")
        elif stripped.startswith("- ") or stripped.startswith("* "):
            doc.add_paragraph(stripped[2:], style="List Bullet")
        # Numbered list
        elif len(stripped) > 2 and stripped[0].isdigit() and stripped[1] in ".、":
            doc.add_paragraph(stripped[2:].lstrip(), style="List Number")
        # Blockquote
        elif stripped.startswith("> "):
            p = doc.add_paragraph(stripped[2:])
            p.paragraph_format.left_indent = Pt(20)
        # Empty line
        elif stripped == "":
            doc.add_paragraph()
        # Plain paragraph (inline bold/italic stripped for simplicity)
        else:
            text = stripped.replace("**", "").replace("*", "").replace("`", "")
            doc.add_paragraph(text)

    if in_table:
        flush_table()

    doc.save(str(docx_path))
    print(f"✅ 已生成 DOCX：{docx_path}")


# ── Feishu upload ─────────────────────────────────────────────────────────────

def upload_file_to_drive(token: str, docx_path: Path, folder_token: str = "") -> str:
    """Upload a file to Feishu Drive and return the file_token."""
    import urllib.request

    file_bytes = docx_path.read_bytes()
    file_size = len(file_bytes)
    file_name = docx_path.name

    # Multipart form upload
    boundary = "----FeishuUploadBoundary"
    body_parts = []

    def field(name, value):
        return (
            f"--{boundary}\r\n"
            f'Content-Disposition: form-data; name="{name}"\r\n\r\n'
            f"{value}\r\n"
        ).encode()

    body_parts.append(field("file_name", file_name))
    body_parts.append(field("parent_type", "explorer"))
    body_parts.append(field("parent_node", folder_token))
    body_parts.append(field("size", str(file_size)))
    body_parts.append(
        f"--{boundary}\r\n"
        f'Content-Disposition: form-data; name="file"; filename="{file_name}"\r\n'
        f"Content-Type: application/vnd.openxmlformats-officedocument.wordprocessingml.document\r\n\r\n"
        .encode()
        + file_bytes
        + b"\r\n"
    )
    body_parts.append(f"--{boundary}--\r\n".encode())

    body = b"".join(body_parts)

    req = urllib.request.Request(
        f"{BASE_URL}/drive/v1/files/upload_all",
        data=body,
        headers={
            "Authorization": f"Bearer {token}",
            "Content-Type": f"multipart/form-data; boundary={boundary}",
        },
    )
    with urllib.request.urlopen(req) as resp:
        data = json.loads(resp.read())

    if data.get("code") != 0:
        print(f"❌ 上传文件失败：{data}")
        sys.exit(1)

    file_token = data["data"]["file_token"]
    print(f"✅ 文件上传成功，file_token：{file_token}")
    return file_token


def import_as_feishu_doc(token: str, file_token: str, file_name: str, folder_token: str = "") -> str:
    """Import a Drive file as a Feishu Doc and return the document URL."""
    import urllib.request

    payload = {
        "file_extension": "docx",
        "file_token": file_token,
        "type": "docx",
        "point": {
            "mount_type": 1,
            "mount_key": folder_token,
        },
    }
    req = urllib.request.Request(
        f"{BASE_URL}/drive/v1/import_tasks",
        data=json.dumps(payload).encode(),
        headers={
            "Authorization": f"Bearer {token}",
            "Content-Type": "application/json",
        },
    )
    with urllib.request.urlopen(req) as resp:
        data = json.loads(resp.read())

    if data.get("code") != 0:
        print(f"❌ 创建导入任务失败：{data}")
        sys.exit(1)

    ticket = data["data"]["ticket"]
    print(f"⏳ 导入任务已创建，ticket：{ticket}，等待完成...")

    # Poll for completion
    for _ in range(20):
        time.sleep(2)
        req2 = urllib.request.Request(
            f"{BASE_URL}/drive/v1/import_tasks/{ticket}",
            headers={"Authorization": f"Bearer {token}"},
        )
        with urllib.request.urlopen(req2) as resp2:
            status_data = json.loads(resp2.read())

        result = status_data.get("data", {}).get("result", {})
        job_status = result.get("job_status")

        if job_status == 0:  # Success
            doc_token = result.get("token")
            doc_type = result.get("type", "docx")
            url = f"https://bytedance.feishu.cn/{doc_type}/{doc_token}"
            print(f"🎉 飞书文档创建成功！")
            print(f"   文档链接：{url}")
            return url
        elif job_status in (1, 2):  # In progress
            print(f"   导入中（状态：{job_status}）...")
        else:
            print(f"❌ 导入失败，状态：{job_status}，详情：{status_data}")
            sys.exit(1)

    print("❌ 导入超时，请手动检查飞书云空间")
    sys.exit(1)


# ── Main ──────────────────────────────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(description="上传招标解析报告到飞书文档")
    parser.add_argument("report_md", help="Markdown 报告文件路径")
    parser.add_argument("--folder-token", default="", help="目标飞书文件夹的 token（留空则存入根目录）")
    args = parser.parse_args()

    md_path = Path(args.report_md)
    if not md_path.exists():
        print(f"❌ 文件不存在：{md_path}")
        sys.exit(1)

    app_id, app_secret = load_credentials()

    # Step 1: Convert Markdown → DOCX
    docx_path = md_path.with_suffix(".docx")
    print(f"📄 正在转换 Markdown → DOCX...")
    markdown_to_docx(md_path, docx_path)

    # Step 2: Auth
    print(f"🔐 获取飞书访问令牌...")
    token = get_tenant_token(app_id, app_secret)

    # Step 3: Upload to Drive
    print(f"☁️  上传到飞书云空间...")
    file_token = upload_file_to_drive(token, docx_path, args.folder_token)

    # Step 4: Import as Feishu Doc
    print(f"📝 转换为飞书文档...")
    url = import_as_feishu_doc(token, file_token, md_path.stem, args.folder_token)

    # Cleanup temp docx
    docx_path.unlink(missing_ok=True)

    return url


if __name__ == "__main__":
    main()
