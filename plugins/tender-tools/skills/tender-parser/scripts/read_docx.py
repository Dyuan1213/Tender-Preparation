#!/usr/bin/env python3
"""
Fallback script to extract text from DOCX files when the Read tool
doesn't capture all content (e.g., tables, headers, footers).

Usage:
    python scripts/read_docx.py <path_to_docx> [--output <output_txt>]
"""

import sys
import argparse
from pathlib import Path


def extract_text(docx_path: str, output_path: str | None = None) -> str:
    try:
        from docx import Document
    except ImportError:
        print("python-docx not installed. Run: pip install python-docx", file=sys.stderr)
        sys.exit(1)

    doc = Document(docx_path)
    sections = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            # Preserve heading structure
            if para.style.name.startswith("Heading"):
                level = para.style.name.replace("Heading ", "")
                prefix = "#" * int(level) if level.isdigit() else "##"
                sections.append(f"{prefix} {text}")
            else:
                sections.append(text)

    # Extract tables
    for i, table in enumerate(doc.tables):
        sections.append(f"\n[表格 {i + 1}]")
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            sections.append(" | ".join(cells))

    full_text = "\n".join(sections)

    if output_path:
        Path(output_path).write_text(full_text, encoding="utf-8")
        print(f"已保存到: {output_path}")
    else:
        print(full_text)

    return full_text


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Extract text from DOCX file")
    parser.add_argument("docx_path", help="Path to the .docx file")
    parser.add_argument("--output", "-o", help="Output text file path (optional)")
    args = parser.parse_args()

    extract_text(args.docx_path, args.output)
